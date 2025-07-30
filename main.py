# main.py

import os
import uvicorn
import asyncio
import json
from datetime import datetime
from fastapi import FastAPI, HTTPException, Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
from dotenv import load_dotenv
import requests
from io import BytesIO
import PyPDF2
import email
import mimetypes
from urllib.parse import urlparse
from pathlib import Path
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import olefile
    import struct
except ImportError:
    olefile = None

# --- LangChain Imports ---
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

# --- 1. Initial Setup and Configuration ---
load_dotenv()
AUTH_TOKEN = "2b55e57dd2584f97b52854b0738dc5608ab353c4fbc8d0409b738b7b21218fbb"

# --- 2. Initialize Models (Done once on startup) ---
print("Initializing models... The first run may take a few minutes to download the embedding model.")

# Using a powerful, retrieval-focused embedding model for high accuracy
embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5")
llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.0)

print("Models initialized successfully.")

# --- 3. FastAPI Application Setup ---
app = FastAPI(title="HackRx 6.0 - Multi-Format Document RAG System", version="1.3")

class QueryRequest(BaseModel):
    documents: str = Field(..., description="A public URL to the document (.pdf, .doc, .docx, .eml, .txt) to be processed.")
    questions: list[str]

class QueryResponse(BaseModel):
    answers: list[str]

bearer_scheme = HTTPBearer()
def validate_token(credentials: HTTPAuthorizationCredentials = Depends(bearer_scheme)):
    if credentials.scheme != "Bearer" or credentials.credentials != AUTH_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid or missing authentication token")

# --- Helper Functions for Document Processing ---
def get_file_extension_from_url(url: str) -> str:
    """Extract file extension from URL"""
    parsed_url = urlparse(url)
    path = Path(parsed_url.path)
    return path.suffix.lower()

def detect_file_type(content: bytes, url: str) -> str:
    """Detect file type based on content and URL"""
    extension = get_file_extension_from_url(url)
    
    # If extension is clear, use it
    if extension in ['.pdf', '.txt', '.eml', '.doc', '.docx']:
        return extension
    
    # Try to detect from content
    if content.startswith(b'%PDF'):
        return '.pdf'
    elif content.startswith(b'PK'):  # ZIP-based formats like .docx
        return '.docx'
    elif content.startswith(b'\xd0\xcf\x11\xe0'):  # OLE format (.doc)
        return '.doc'
    elif b'Content-Type:' in content[:1000] or b'From:' in content[:1000]:
        return '.eml'
    else:
        # Default to text if can't determine
        return '.txt'

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content"""
    pdf_bytes = BytesIO(content)
    pdf_reader = PyPDF2.PdfReader(pdf_bytes)
    text = "\n".join(page.extract_text() for page in pdf_reader.pages)
    return text

def extract_text_from_txt(content: bytes) -> str:
    """Extract text from TXT content"""
    try:
        # Try UTF-8 first
        return content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Try latin-1 as fallback
            return content.decode('latin-1')
        except UnicodeDecodeError:
            # Try cp1252 as last resort
            return content.decode('cp1252', errors='ignore')

def extract_text_from_eml(content: bytes) -> str:
    """Extract text from EML (email) content"""
    try:
        # Parse the email
        email_message = email.message_from_bytes(content)
        
        text_parts = []
        
        # Extract subject
        subject = email_message.get('Subject', '')
        if subject:
            text_parts.append(f"Subject: {subject}")
        
        # Extract sender and recipient info
        from_addr = email_message.get('From', '')
        if from_addr:
            text_parts.append(f"From: {from_addr}")
        
        to_addr = email_message.get('To', '')
        if to_addr:
            text_parts.append(f"To: {to_addr}")
        
        date = email_message.get('Date', '')
        if date:
            text_parts.append(f"Date: {date}")
        
        # Extract body content
        if email_message.is_multipart():
            for part in email_message.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    body = part.get_payload(decode=True)
                    if body:
                        text_parts.append(body.decode('utf-8', errors='ignore'))
                elif content_type == "text/html":
                    # Basic HTML stripping for HTML parts
                    html_body = part.get_payload(decode=True)
                    if html_body:
                        import re
                        html_text = html_body.decode('utf-8', errors='ignore')
                        # Remove HTML tags
                        clean_text = re.sub('<[^<]+?>', '', html_text)
                        text_parts.append(clean_text)
        else:
            body = email_message.get_payload(decode=True)
            if body:
                text_parts.append(body.decode('utf-8', errors='ignore'))
        
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error parsing EML file: {e}")
        # Fallback to treating as plain text
        return extract_text_from_txt(content)

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content"""
    if DocxDocument is None:
        raise HTTPException(status_code=500, detail="python-docx library not installed. Cannot process .docx files.")
    
    try:
        docx_bytes = BytesIO(content)
        doc = DocxDocument(docx_bytes)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {e}")

def extract_text_from_doc(content: bytes) -> str:
    """Extract text from DOC content (legacy Word format)"""
    # This is a basic implementation - for production use, consider using python-docx2txt or similar
    try:
        # Try to extract readable text from the binary DOC format
        # This is a simplified approach and may not work for all DOC files
        text = content.decode('latin-1', errors='ignore')
        
        # Remove control characters and clean up
        import re
        # Remove null bytes and control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', ' ', text)
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Extract readable words (simple heuristic)
        words = []
        for word in text.split():
            # Keep words that are mostly alphabetic
            if len(word) > 2 and sum(c.isalpha() for c in word) / len(word) > 0.7:
                words.append(word)
        
        if not words:
            raise Exception("No readable text found")
        
        return " ".join(words)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOC file: {e}. Consider converting to DOCX or PDF format.")

def extract_text_from_document(content: bytes, url: str) -> str:
    """Main function to extract text from any supported document type"""
    file_type = detect_file_type(content, url)
    
    print(f"Detected file type: {file_type}")
    
    if file_type == '.pdf':
        return extract_text_from_pdf(content)
    elif file_type == '.txt':
        return extract_text_from_txt(content)
    elif file_type == '.eml':
        return extract_text_from_eml(content)
    elif file_type == '.docx':
        return extract_text_from_docx(content)
    elif file_type == '.doc':
        return extract_text_from_doc(content)
    else:
        # Default to text extraction
        return extract_text_from_txt(content)

# --- Helper Functions for File Storage ---
def create_storage_directories():
    """Create necessary storage directories"""
    base_dir = Path(".")
    documents_dir = base_dir / "documents"
    qa_dir = base_dir / "questionandanswers"
    
    documents_dir.mkdir(exist_ok=True)
    qa_dir.mkdir(exist_ok=True)
    
    return documents_dir, qa_dir

def get_filename_from_url(url: str) -> str:
    """Extract filename from URL or generate one"""
    parsed_url = urlparse(url)
    path = Path(parsed_url.path)
    
    if path.name:
        return path.name
    else:
        # Generate filename based on timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        extension = get_file_extension_from_url(url) or '.txt'
        return f"document_{timestamp}{extension}"

def save_document(content: bytes, extracted_text: str, url: str, documents_dir: Path) -> str:
    """Save both original document and extracted text"""
    try:
        filename = get_filename_from_url(url)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create a unique filename with timestamp
        name, ext = os.path.splitext(filename)
        unique_filename = f"{name}_{timestamp}{ext}"
        
        # Save original document
        original_path = documents_dir / unique_filename
        with open(original_path, 'wb') as f:
            f.write(content)
        
        # Save extracted text
        text_filename = f"{name}_{timestamp}_extracted.txt"
        text_path = documents_dir / text_filename
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(f"Source URL: {url}\n")
            f.write(f"Extraction Date: {datetime.now().isoformat()}\n")
            f.write(f"Original Filename: {filename}\n")
            f.write("-" * 50 + "\n")
            f.write(extracted_text)
        
        print(f"Document saved: {original_path}")
        print(f"Extracted text saved: {text_path}")
        
        return unique_filename
    except Exception as e:
        print(f"Error saving document: {e}")
        return filename

def save_qa_results(questions: list[str], answers: list[str], context_chunks: list, url: str, document_filename: str, qa_dir: Path):
    """Save questions, answers, and context to JSON file"""
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        qa_filename = f"qa_session_{timestamp}.json"
        qa_path = qa_dir / qa_filename
        
        qa_data = {
            "session_info": {
                "timestamp": datetime.now().isoformat(),
                "source_url": url,
                "document_filename": document_filename,
                "total_questions": len(questions)
            },
            "qa_pairs": []
        }
        
        for i, (question, answer) in enumerate(zip(questions, answers)):
            qa_pair = {
                "question_id": i + 1,
                "question": question,
                "answer": answer,
                "context_used": [chunk.page_content for chunk in context_chunks] if context_chunks else []
            }
            qa_data["qa_pairs"].append(qa_pair)
        
        with open(qa_path, 'w', encoding='utf-8') as f:
            json.dump(qa_data, f, indent=2, ensure_ascii=False)
        
        print(f"Q&A results saved: {qa_path}")
        
    except Exception as e:
        print(f"Error saving Q&A results: {e}")

# --- 4. The RAG Logic (Optimized for Speed and Accuracy) ---
async def process_query_with_rag(doc_url: str, questions: list[str]) -> list[str]:
    # Step 0: Create storage directories
    documents_dir, qa_dir = create_storage_directories()
    
    # Step 1: Data Ingestion from any public URL
    print(f"Fetching document from: {doc_url}")
    try:
        response = requests.get(doc_url)
        response.raise_for_status() # Raise an exception for bad status codes (4xx or 5xx)
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=400, detail=f"Failed to download document: {e}")

    # Step 1.5: Extract text based on document type
    try:
        text = extract_text_from_document(response.content, doc_url)
        print(f"Extracted {len(text)} characters from the document.")
        
        if len(text.strip()) == 0:
            raise HTTPException(status_code=400, detail="No text content could be extracted from the document.")
        
        # Save the document and extracted text
        document_filename = save_document(response.content, text, doc_url, documents_dir)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from document: {e}")

    # Step 2: Chunking
    docs = [Document(page_content=text)]
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=300)
    documents = text_splitter.split_documents(docs)

    # Step 3: Vector Store Creation
    print("Creating FAISS vector store with upgraded embeddings...")
    vectorstoredb = FAISS.from_documents(documents=documents, embedding=embeddings)
    retriever = vectorstoredb.as_retriever(search_kwargs={"k": 10})
    print("Vector store created.")

    # Step 4: Prompt and Chain Definition
    prompt = ChatPromptTemplate.from_template(
       """
        You are a highly accurate question-answering system based on the provided document context. Your task is to extract the answer from the context and give the output in a concise and professional manner.
        Please refer the document context properly and answer the question based on the context provided. Make sure to follow these instructions carefully.
        **Context:**
        ---
        {context}
        ---
        **Question:** {input}
        **Instructions:**
        1. Your answer must be extracted from the context.
        2. Do not just copy-paste from the context.
        3. Synthesize the relevant information into a complete, professional sentence or two.
        4. Do not provide any additional information or explanations.
        5. The answer should be concise and directly address the question, should not exceed 50 words per answer.
        6. If the answer is ambiguos, then search again in the context and provide the most relevant answer.
        7. If the answer is not present in the context, respond with "This information is not available in the provided document context."
        8. For questions that require particular context, take the context as a reference and provide the answer based on that context.
        9. When the question is specific about something, make sure to provide the answer based on the specific context.
        **Answer:**
        """
    )
    document_chain = create_stuff_documents_chain(llm=llm, prompt=prompt)
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    # Step 5: Process all questions concurrently for speed
    print(f"Processing {len(questions)} questions concurrently...")
    tasks = [retrieval_chain.ainvoke({"input": q}) for q in questions]
    results = await asyncio.gather(*tasks)
    answers = [res['answer'] for res in results]
    
    # Step 6: Save Q&A results with context
    try:
        # Get context from the first result for reference (all should have similar chunks)
        context_chunks = results[0].get('context', []) if results else []
        save_qa_results(questions, answers, context_chunks, doc_url, document_filename, qa_dir)
    except Exception as e:
        print(f"Error saving Q&A results: {e}")
    
    print("All questions processed.")
    return answers

# --- 5. API Endpoint ---
@app.post(
    "/api/v1/hackrx/run",
    response_model=QueryResponse,
    dependencies=[Depends(validate_token)],
    tags=["Query System"]
)
async def run_submission(request: QueryRequest):
    try:
        final_answers = await process_query_with_rag(request.documents, request.questions)
        return QueryResponse(answers=final_answers)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- 6. To run the app ---
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)